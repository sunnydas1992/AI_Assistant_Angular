"""
Document Processor - Extract text from PDF, DOCX, TXT, MD. File-like: .name, .type, .getvalue().
"""

import hashlib
import logging
from io import BytesIO
from typing import List, Dict, Any, Optional

import pypdf
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config.settings import RAGConfig

logger = logging.getLogger(__name__)


def extract_text_for_chat_attachment(name: str, data: bytes, content_type: str = "") -> tuple[str, bool]:
    """
    Extract human-readable text for chat context (PDF, DOCX, plain text).

    Returns:
        (text, used_extractor): If used_extractor is True, text is from a structured
        extractor (or a placeholder if none). If False, caller should decode bytes as UTF-8.
    """
    if not isinstance(data, bytes):
        return (str(data), True)
    file_name = (name or "").lower()
    ct = (content_type or "").lower()

    def _is_pdf() -> bool:
        return ct == "application/pdf" or file_name.endswith(".pdf")

    def _is_docx() -> bool:
        return (
            ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or file_name.endswith(".docx")
        )

    def _is_plain() -> bool:
        return ct in ("text/plain", "text/markdown") or file_name.endswith((".txt", ".md"))

    try:
        if _is_pdf():
            pdf_reader = pypdf.PdfReader(BytesIO(data))
            text = "".join(page.extract_text() or "" for page in pdf_reader.pages)
            if not text.strip():
                return (
                    "[PDF: no extractable text (it may be scanned or image-only). "
                    "Try OCR, export as text, or paste the content.]",
                    True,
                )
            return (text, True)
        if _is_docx():
            doc = Document(BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs)
            if not text.strip():
                return ("[DOCX: no extractable text.]", True)
            return (text, True)
        if file_name.endswith(".doc") and not _is_docx():
            return (
                "[Legacy .doc is not supported; save as .docx or PDF and attach again.]",
                True,
            )
        if _is_plain():
            return (data.decode("utf-8", errors="replace"), True)
    except Exception as exc:
        logger.warning("Chat attachment extract failed for '%s': %s", name, exc)
        return (f"[Could not read file: {exc}]", True)

    return ("", False)


class DocumentProcessor:
    """Process uploaded files (expects file-like with .name, .type, .getvalue())."""

    SUPPORTED_TYPES = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "text",
        "text/markdown": "text",
    }

    def __init__(self, config: Optional[RAGConfig] = None):
        self.config = config or RAGConfig()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap
        )

    def process_uploaded_files(self, uploaded_files) -> List[Dict[str, Any]]:
        documents = []
        if not uploaded_files:
            return documents
        for file in uploaded_files:
            doc = self._process_single_file(file)
            if doc:
                documents.append(doc)
        return documents

    def _process_single_file(self, file) -> Optional[Dict[str, Any]]:
        source_title = getattr(file, 'name', None) or getattr(file, 'filename', 'unknown')
        try:
            text = self._extract_text(file)
            if text is None:
                logger.warning("Unsupported or unrecognised file type for '%s', skipping", source_title)
                return None
            if not text.strip():
                logger.warning("No extractable text in '%s' (scanned / image-only PDF?), skipping", source_title)
                return None
            content = f"Source: Uploaded File '{source_title}'\n\n{text}"
            return {
                "text": content,
                "meta": {"source_type": "upload", "title": source_title, "url": ""}
            }
        except Exception as exc:
            logger.warning("Failed to process file '%s': %s", source_title, exc)
            return None

    def _extract_text(self, file) -> Optional[str]:
        file_type = getattr(file, 'type', None) or getattr(file, 'content_type', '')
        file_name = (getattr(file, 'name', '') or getattr(file, 'filename', '')).lower()
        if file_type == "application/pdf" or file_name.endswith(".pdf"):
            return self._extract_pdf_text(file)
        if (file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or file_name.endswith(".docx")):
            return self._extract_docx_text(file)
        if (file_type in ("text/plain", "text/markdown")
                or file_name.endswith((".txt", ".md"))):
            return self._extract_plain_text(file)
        return None

    def _extract_pdf_text(self, file) -> str:
        data = file.getvalue() if hasattr(file, 'getvalue') else file.read()
        pdf_reader = pypdf.PdfReader(BytesIO(data))
        return "".join(page.extract_text() or "" for page in pdf_reader.pages)

    def _extract_docx_text(self, file) -> str:
        data = file.getvalue() if hasattr(file, 'getvalue') else file.read()
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    def _extract_plain_text(self, file) -> str:
        data = file.getvalue() if hasattr(file, 'getvalue') else file.read()
        if isinstance(data, bytes):
            return data.decode("utf-8", errors="replace")
        return data

    def split_text(self, text: str) -> List[str]:
        return self.text_splitter.split_text(text)

    def generate_chunk_id(self, title: str, index: int, chunk: str) -> str:
        base = f"{title}::{index}::{chunk}"
        return hashlib.sha256(base.encode("utf-8")).hexdigest()[:24]

    def prepare_chunks_for_storage(self, documents: List[Dict[str, Any]]) -> tuple:
        all_chunks, all_metas, all_ids = [], [], []
        for doc in documents:
            text = doc["text"]
            meta = doc.get("meta", {})
            chunks = self.split_text(text)
            for idx, chunk in enumerate(chunks):
                chunk_id = self.generate_chunk_id(meta.get('title', ''), idx, chunk)
                chunk_meta = dict(meta)
                chunk_meta["chunk_index"] = idx
                all_chunks.append(chunk)
                all_metas.append(chunk_meta)
                all_ids.append(chunk_id)
        return all_chunks, all_metas, all_ids

    @staticmethod
    def get_supported_extensions() -> List[str]:
        return ["pdf", "docx", "txt", "md"]
